import tkinter as tk
from tkinter import*
from tkinter import ttk
from tkinter import messagebox as mb
from openpyxl import load_workbook
from FullScreen import FullScreenApp
from time import strftime
from CustomerManagement import AddCustomerWindow
from Inventory import NewProduct,SearchWindow
from docx import Document
import random
from datetime import datetime
import hashlib
class MainSession(tk.Tk):
    def __init__(self,account,position,username):
        global tree,canvas,total,price_list,barcode_entry,cart,category_list,name_list,category_combobox,name_list,name_combobox,price,quantity,phone,name,loyalty_program,customer_entry,total_unformat,barcode,discount,pay_amount
        global delivery_fee,delivery_fee_entry,money_received,change_amount,money_received_entry,clock,name,phone,shipping_status,discount_entry,cart,account_username
        # create self
        tk.Tk.__init__(self)
        FullScreenApp(self)
        self.title("Bán hàng")
        # background
        canvas=Canvas(master=self,width=1920,height=1080,bg="#000000")
        bg=Label(master=self)
        bg_img=PhotoImage(file="main_session.png",master=canvas)
        bg.configure(image=bg_img)
        bg.place(relx=0,rely=0,width=1920,height=1080)
        # buttons
        add_button_img=PhotoImage(master=canvas,file="addbutton_mainsession.png")
        add_barcode=Button(master=self,image=add_button_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.SearchBarcode)
        add_barcode.place(x=390,y=350)

        add_customer=Button(master=self,image=add_button_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.SearchCustomer)
        add_customer.place(x=1820,y=127)

        log_out_img=PhotoImage(master=canvas,file="logout_mainsession.png")
        log_out_button=Button(master=self,image=log_out_img,command=self.Log_out,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0")
        log_out_button.place(x=41,y=211)

        change_password_img=PhotoImage(master=canvas,file="changepassword_mainsession.png")
        change_password_button=Button(master=self,image=change_password_img,command=ChangePassword,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0")
        change_password_button.place(x=255,y=211)

        add_to_cart_img=PhotoImage(master=canvas,file="addtocart_mainsession.png")
        add_to_cart_button=Button(master=self,image=add_to_cart_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.AddtoCart)
        add_to_cart_button.place(x=41,y=599)

        del_order_img=PhotoImage(master=canvas,file="deletecart_mainsession.png")
        del_order_button=Button(master=self,image=del_order_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.Del_order)
        del_order_button.place(x=255,y=599)

        add_new_product_img=PhotoImage(master=canvas,file="newproduct_mainsession.png")
        add_new_product_button=Button(master=self,image=add_new_product_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=NewProduct)
        add_new_product_button.place(x=41,y=737)

        edit_product_img=PhotoImage(master=canvas,file="editproduct_mainsession.png")
        edit_product_button=Button(master=self,image=edit_product_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=lambda:SearchWindow("Edit"))
        edit_product_button.place(x=255,y=737)

        add_delivery_img=PhotoImage(master=canvas,file="delivery_mainsession.png")
        add_delivery_button=Button(master=self,image=add_delivery_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.AddDeliveryAddress)
        add_delivery_button.place(x=1471,y=737)

        pay_img=PhotoImage(master=canvas,file="pay_mainsession.png")
        pay_button=Button(master=self,image=pay_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.Pay)
        pay_button.place(x=1472,y=891)

        account_name = StringVar(self)
        account_name.set(account)
        account_name_entry = Entry(self, relief="flat", font="Montserrat", state="readonly", textvariable=account_name,
                            readonlybackground="#ffffff", cursor="hand2")
        account_name_entry.place(x=101, y=171, width=346, height=19)
        account_username=username


        category_combobox=ttk.Combobox(master=self,state="readonly")
        category_combobox.place(x=207,y=427,width=242,height=23)
        category_combobox.bind("<<ComboboxSelected>>",self.SelectCategory)

        name_combobox = ttk.Combobox(master=self, state="readonly")
        name_combobox.bind("<<ComboboxSelected>>",self.SelectName)
        name_combobox.place(x=207, y=470, width=242, height=23)

        price=StringVar(self)
        price_entry=Entry(self, relief="flat", font="Montserrat", state="readonly", textvariable=price,
                           readonlybackground="#ffffff", cursor="hand2")
        price_entry.place(x=207,y=513,width=242,height=23)

        quantity=StringVar(self)
        quantity_entry=Entry(self,relief="flat",font="Montserrat",textvariable=quantity,cursor="hand2")
        quantity_entry.place(x=207,y=556,width=242,height=23)

        barcode=StringVar(self)
        barcode_entry=Entry(self,relief="flat",font="Montserrat",textvariable=barcode,cursor="hand2")
        barcode_entry.place(x=44,y=353,width=324,height=20)

        customer_phone=StringVar(self)
        customer_entry = Entry(self, relief="flat", font="Montserrat", textvariable=customer_phone, cursor="hand2")
        customer_entry.place(x=1474, y=130, width=324, height=20)

        name = StringVar(self)
        name_entry = Entry(self, relief="flat", font="Montserrat", state="readonly", textvariable=name,
                           readonlybackground="#ffffff", cursor="hand2")
        name_entry.place(x=1637, y=169, width=242, height=23)
        name.set("Khách lẻ")

        phone = StringVar(self)
        phone_entry = Entry(self, relief="flat", font="Montserrat", state="readonly", textvariable=phone,
                           readonlybackground="#ffffff", cursor="hand2")
        phone_entry.place(x=1637, y=212, width=242, height=23)

        loyalty_program = StringVar(self)
        loyalty_program_entry = Entry(self, relief="flat", font="Montserrat", state="readonly", textvariable=loyalty_program,
                           readonlybackground="#ffffff", cursor="hand2")
        loyalty_program_entry.place(x=1637, y=255, width=242, height=23)

        total = StringVar(self)
        total.set("0")
        total_entry = Entry(self, relief="flat", font="Montserrat", state="readonly", textvariable=total,
                           readonlybackground="#ffffff", cursor="hand2")
        total_entry.place(x=1637, y=333, width=242, height=23)
        total_unformat=0
        # reg_total=self.register(self.CalculateTotal)
        # total_entry.config(validate="key",validatecommand=(reg_total,"%P"))

        discount=StringVar(self)
        discount.set("0")
        discount_entry= Entry(self, relief="flat", font="Montserrat", textvariable=discount, cursor="hand2")
        discount_entry.place(x=1637, y=376, width=164, height=23)
        apply_discount_img=PhotoImage(master=canvas,file="addbutton_mainsession.png")
        apply_discount=Button(master=self,image=apply_discount_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.Discount)
        apply_discount.place(x=1820,y=376)
        discount.trace("w",self.Format_Discount)

        delivery_fee = StringVar(self)
        delivery_fee.set("0")
        delivery_fee_entry = Entry(self, relief="flat", font="Montserrat", textvariable=delivery_fee,
                            readonlybackground="#ffffff", cursor="hand2")
        delivery_fee_entry.place(x=1637, y=419, width=242, height=23)
        delivery_fee.trace("w",self.Format_Delivery_Fee)
        add_delivery_fee_img=PhotoImage(master=canvas,file="addbutton_mainsession.png")
        add_delivery_fee_button=Button(master=self,image=add_delivery_fee_img,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.AddDeliveryFee)
        add_delivery_fee_button.place(x=1820,y=419)

        pay_amount = StringVar(self)
        pay_amount.set("0")
        pay_amount_entry = Entry(self,relief="flat", font="Montserrat",state="readonly", textvariable=pay_amount,
                                   readonlybackground="#ffffff", cursor="hand2")
        pay_amount_entry.place(x=1650, y=497, width=227, height=38)

        money_received = StringVar(self)
        money_received.set("0")
        money_received_entry = Entry(self, relief="flat", font="Montserrat", textvariable=money_received,
                                 readonlybackground="#ffffff", cursor="hand2")
        money_received.trace("w",self.Format_Received_Money)
        money_received_entry.place(x=1650, y=559, width=227, height=19)

        change_amount = StringVar(self)
        change_amount_entry = Entry(self, relief="flat", font="Montserrat", textvariable=change_amount,state="readonly",
                                     readonlybackground="#ffffff", cursor="hand2")
        change_amount_entry.place(x=1650, y=602, width=227, height=19)

        payment_combobox=ttk.Combobox(master=self, state="readonly",value=("Tiền mặt","Thẻ ngân hàng","Chuyển khoản","COD"))
        payment_combobox.place(x=1647, y=643, width=232, height=23)

        clock = Label(self, background="#ffffff", font="Montserrat",cursor="hand2")
        clock.place(x=154,y=930,width=227,height=38)

        shipping_status=False
        self.time()
        # cart
        TableMargin = Frame(self, width=942)
        TableMargin.place(x=484, y=76)
        scrollbary = Scrollbar(TableMargin, orient=VERTICAL)
        tree = ttk.Treeview(TableMargin,
                            columns=("Name", "Price", "Quantity", "Total"),
                            height=46
                            , selectmode="extended", yscrollcommand=scrollbary.set)
        scrollbary.config(command=tree.yview)
        scrollbary.pack(side=RIGHT, fill=Y)
        tree.heading('Name', text="Tên sản phẩm")
        tree.heading('Price', text="Đơn giá")
        tree.heading('Quantity', text="Số lượng")
        tree.heading('Total', text="Thành tiền")
        tree.column('#0', stretch=NO, minwidth=0, width=0)
        tree.column('#1', stretch=NO, minwidth=0, width=300)
        tree.column('#2', stretch=NO, minwidth=0, width=240)
        tree.column('#3', stretch=NO, minwidth=0, width=150)
        tree.column('#4', stretch=NO, minwidth=0, width=240)
        tree.pack()
        tree.bind("<Double-1>",self.EditCart)
        cart={}
        category_list=[]
        name_list={}
        excelfile=load_workbook("ProductsList.xlsx")
        for row in excelfile.worksheets[0]:
            if row[1].value in category_list:
                pass
            else:
                category_list.append(row[1].value)
                name_list[row[1].value]=[]
        for row in excelfile.worksheets[0]:
            if row[0].value in name_list[row[1].value]:
                pass
            else:
                name_list[row[1].value].append(row[0].value)
        name_list.pop("Category")
        category_list.remove("Category")
        category_combobox.configure(values=tuple(category_list))
        excelfile.close()
        self.position=position
        self.mainloop()
    def Format_Delivery_Fee(*events):
        cache = delivery_fee.get()
        cache = cache.replace(",", '')
        try:
            cache = int(cache)
            cache = format(cache, ',')
        except:
            if cache=="":
                pass
            else:
                mb.showerror("Thông báo","Phí giao hàng không hợp lệ")
        delivery_fee_entry.delete(0,END)
        delivery_fee_entry.insert(END,cache)
    def Format_Received_Money(*events):
        cache = money_received.get()
        cache = cache.replace(",", '')
        pay_amount_value=pay_amount.get().replace(",","")
        try:
            cache = int(cache)
            if cache >= int(pay_amount_value):
                delivery=delivery_fee.get()
                delivery=delivery.replace(",","")
                change_amount.set(str("{:,}".format(cache-int(pay_amount_value)-int(delivery))))
            else:
                change_amount.set("Tiền khách đưa không hợp lệ")
            cache = format(cache, ',')
        except:
            if cache=="":
                change_amount.set("0")
            else:
                mb.showerror("Thông báo","Tiền khách đưa không hợp lệ")
        money_received_entry.delete(0,END)
        money_received_entry.insert(END,cache)
    def Log_out(self):
        self.destroy()
    def Format_Discount(*events):
        cache = discount.get()
        cache = cache.replace(",", '')
        try:
            cache = int(cache)
            cache = format(cache, ',')
        except:
            if cache == "" or "%" in cache:
                pass
            else:
                mb.showerror("Thông báo", "Phí giao hàng không hợp lệ")
        discount_entry.delete(0,END)
        discount_entry.insert(END,cache)
    def AddDeliveryFee(self):
        input=delivery_fee.get()
        try:
            input=input.replace(",","")
            pay_amount.set(str("{:,}".format(eval(str(total_unformat)+"+"+input))))
        except:
            if input=="":
                pay_amount.set(str("{:,}".format(total_unformat)))
                pass
            else:
                mb.showerror("Thông báo", "Phí giao hàng không hợp lệ")
    def AddDeliveryAddress(self):
        Delivery_window()
    def EditCart(self,event):
        item = tree.selection()[0]
        name=tree.item(item)["values"][0]
        category_combobox.current(category_combobox["values"].index(cart[name]["category"]))
        name_combobox.current(name_combobox["values"].index(cart[name]["name"]))
        quantity.set(cart[name]["quantity"])
        price.set(str("{:,}".format(int(cart[name]["price"]))))
    def Discount(self):
        entry=discount.get()
        entry=entry.replace(",","")
        percentage=False
        if "%" in entry:
            percentage=True
        if percentage==True:
            if float(entry.removesuffix("%")) > 0 and float(entry.removesuffix("%")) <= 100:
                pay_amount.set(str("{:,}".format(round(eval(str(total_unformat)+"*"+"(1-"+entry.removesuffix("%")+"/100)")))))
            else:
                mb.showerror("Thông báo","Chiết khấu tối đa 100%")
        else:
            if entry.isdigit() == True:
                if int(entry) > total_unformat:
                    mb.showerror("Thông báo","Chiết khấu của bạn lớn hơn giá trị đơn hàng")
                else:
                    pay_amount.set(str("{:,}".format(eval(str(total_unformat)+"-"+entry))))
            else:
                mb.showerror("Thông báo","Bạn nhập sai giá trị.\nCú pháp: Chiết khấu theo phần trăm ('xxx%'); Chiết khấu theo giá trị cụ thể ('xxxxxx')")
    def SearchBarcode(self):
        excelfile=load_workbook("ProductsList.xlsx")
        for row in excelfile.worksheets[0]:
            if barcode_entry.get()==row[6].value:
                category_combobox.current(category_combobox["values"].index(row[1].value))
                self.SelectCategory()
                name_combobox.current(name_combobox["values"].index(row[0].value))
                price.set(f"{row[4].value:,}")
                self.price_cache=str(row[4].value)
                quantity.set("1")
                break
        excelfile.close()
    def SelectCategory(self,event=None):
        name_combobox.configure(values=tuple(name_list[category_combobox.get()]))
        name_combobox.current(0)
    def SelectName(self,event=None):
        if event:
            excelfile=load_workbook("ProductsList.xlsx")
            for row in excelfile.worksheets[0]:
                if name_combobox.get()==row[0].value:
                    price.set(f"{row[4].value:,}")
                    self.price_cache = str(row[4].value)
                    quantity.set("1")
                    break
            excelfile.close()
    def AddtoCart(self):
        global total,total_unformat,cart
        excelfile=load_workbook("ProductsList.xlsx")
        if name_combobox.get() in cart:
            if quantity.get()=="0":
                total_unformat=total_unformat-eval(cart[name_combobox.get()]["quantity"]+"*"+cart[name_combobox.get()]["price"])
                total.set(str("{:,}".format(total_unformat)))
                pay_amount.set(total.get())
                cart.pop(name_combobox.get())
                for child in tree.get_children():
                    if name_combobox.get() in tree.item(child)["values"]:
                        tree.delete(child)
                        break
                barcode.set("")
                pass
            else:
                for child in tree.get_children():
                    if name_combobox.get() in tree.item(child)["values"]:
                        for row in excelfile.worksheets[0]:
                            if row[0].value==name_combobox.get():
                                if eval(str(quantity.get())+"+"+str(tree.item(child)["values"][2]))>float(row[2].value):
                                    mb.showerror("Thông báo","Quá số lượng tồn kho")
                                else:
                                    cart[name_combobox.get()]["quantity"] = str(eval(str(quantity.get()) + "+" + str(tree.item(child)["values"][2])))
                                    total_unformat=total_unformat+eval(quantity.get()+"*"+cart[name_combobox.get()]["price"])
                                    tree.delete(child)
                                    tree.insert("",0,values=(name_combobox.get(),self.price_cache,cart[name_combobox.get()]["quantity"],str("{:,}".format(eval(cart[name_combobox.get()]["quantity"]+"*"+cart[name_combobox.get()]["price"])))))
                                    total.set(str("{:,}".format(total_unformat)))
                                    pay_amount.set(total.get())
                                    barcode.set("")
                                break
                        break
        else:
            for row in excelfile.worksheets[0]:
                if row[0].value==name_combobox.get():
                    if float(quantity.get())>float(row[2].value):
                        mb.showerror("Thông báo","Quá số lượng tồn kho")
                    else:
                        cart[name_combobox.get()]={"name":name_combobox.get(),"price":self.price_cache,"quantity":str(quantity.get()),"category":category_combobox.get()}
                        total_cache=StringVar(self)
                        total_cache.set(str(eval(self.price_cache+"*"+quantity.get())))
                        total_unformat=total_unformat+int(total_cache.get())
                        tree.insert("",0,values=(name_combobox.get(),"{:,}".format(int(self.price_cache)),str(quantity.get()),str("{:,}".format(int(total_cache.get())))))
                        total.set(str("{:,}".format(total_unformat)))
                        pay_amount.set(total.get())
                        barcode.set("")
                    break
        excelfile.close()
    def time(self):
        time_info = strftime("%H:%M:%S %p")
        try:
            clock.configure(text=time_info)
            clock.after(1000, self.time)
        except:
            pass
    def Del_order(self):
        global total_unformat,price_cache
        tree.delete(*tree.get_children())
        total.set("")
        barcode.set("")
        name.set("")
        phone.set("")
        loyalty_program.set("")
        total_unformat=0
        price_cache=0
    def SearchCustomer(self):
        status=False
        excelfile=load_workbook("CustomerManagement.xlsx")
        for row in excelfile.worksheets[0]:
            if customer_entry.get()==row[4].value:
                name.set(row[0].value)
                phone.set(row[4].value)
                loyalty_program.set(row[6].value)
                status=True
                customer_entry.delete(0,END)
                break
        if status==False:
            add_customer=mb.askyesno("Thông báo","Khách hàng chưa có trong hệ thống. Bạn có muốn thêm khách hàng mới không?")
            if add_customer==True:
                AddCustomerWindow(self.position)
            if add_customer==False:
                customer_entry.delete(0,END)
    def Pay(self):
        profile_store = load_workbook("Profile_store.xlsx")
        data = profile_store["Sheet1"]
        profile = {"name": data["A2"].value, "address": data["B2"].value, "phone": data["C2"].value}
        profile_store.close()
        bill = Document()
        bill.add_heading('Hóa đơn bán hàng', 0)
        bill.add_heading(profile["name"] + "          " + profile["address"] + "          " + profile["phone"], 2)
        ref="OR"+str(random.randint(100000,999999))
        time=datetime.now()
        time=time.strftime("%H:%M:%S %p")
        bill.add_heading("Ref: "+ref+"          "+time,2)
        bill.add_paragraph("")
        customer_info=bill.add_paragraph("Tên khách hàng: "+name.get()+"          "+"Di động: "+phone.get() +"          "+"Thành viên: "+loyalty_program.get())
        # customer_info.add_run("Tên khách hàng:").bold=True
        # customer_info.add_run("Di động:").bold=True
        # customer_info.add_run("Thành viên:").bold = True
        self.cart=bill.add_table(rows=1,cols=4)
        self.cart.style="Medium Shading 1 Accent 1"
        row=self.cart.rows[0].cells
        row[0].text="Tên sản phẩm"
        row[1].text="Đơn giá"
        row[2].text="Số lượng"
        row[3].text="Thành tiền"
        products_excel=load_workbook("ProductsList.xlsx")
        for child in tree.get_children():
            row=self.cart.add_row().cells
            index=0
            for i in tree.item(child)["values"]:
                row[index].text=str(i)
                index+=1
            for row_excel in products_excel.worksheets[0]:
                if row_excel[0].value==row[0]:
                    cart[row[0]]["barcode"]=row_excel[6].value
                    value=float(row_excel[2].value)-float(row[2].text)
                    try:
                        row[2].value=int(value)
                    except:
                        row[2].value=float(value)
                    break
        products_excel.save("ProductsList.xlsx")
        products_excel.close()
        bill.add_paragraph("")
        if shipping_status==True:
            shipping=bill.add_paragraph("Địa chỉ giao hàng: "+address_shipping+"\n"+"Số điện thoại người nhận: "+phone_shipping)
            # shipping.add_run("Địa chỉ giao hàng:").bold=True
            # shipping.add_run("Số điện thoại người nhận:").bold=True
        else:
            pass
        bill.add_paragraph("")
        filename=ref+".docx"
        bill.save(filename)
        excelfile=load_workbook("TransactionManagement.xlsx")
        data=excelfile.worksheets[0]
        time = datetime.now()
        time = time.strftime("%d/%m/%Y %H:%M:%S")
        info=[ref,"Sell: "+ref,pay_amount.get(),"IN",time,phone.get()]
        data.insert_rows(idx=2,amount=1)
        index=0
        for i in ["A2","B2","C2","D2","E2","F2"]:
            data[i].value=info[index]
            index+=1
        excelfile.save("TransactionManagement.xlsx")
        excelfile.close()
        CustomerManagement_excel=load_workbook("CustomerManagement.xlsx")
        for row in CustomerManagement_excel.worksheets[0]:
            if row[4].value==phone.get():
                pay_amount_value=pay_amount.get().replace(",","")
                row[7].value=row[7].value+int(pay_amount_value)
                row[7].number_format =  '#,##0'
                break
        self.Del_order()
        CustomerManagement_excel.save("CustomerManagement.xlsx")
        CustomerManagement_excel.close()
        mb.showinfo("Thông báo","Đơn hàng lưu thành công",master=self)
class Delivery_window(tk.Tk):
    def __init__(self):
        global address_entry, phone_entry_shipping
        tk.Tk.__init__(self)
        self.geometry("400x200")
        self.resizable(0, 0)
        self.title("Thêm địa chỉ giao hàng")
        canvas=Canvas(master=self,width=400,height=200)
        bg_img=PhotoImage(master=canvas,file="delivery.png")
        bg=Label(self,image=bg_img)
        bg.place(x=0,y=0)
        button_img=PhotoImage(master=canvas,file="Save_delivery.png")
        save_button=Button(master=self,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.CheckValid)
        save_button.configure(image=button_img)
        save_button.place(x=134,y=150)
        address_entry=Entry(self,font="Montserrat",relief="flat")
        address_entry.place(x=97,y=62,width=283,height=25)
        phone_entry_shipping = Entry(self, font="Montserrat", relief="flat")
        phone_entry_shipping.place(x=97, y=99, width=283, height=25)
        phone_entry_shipping.insert(0,phone.get())
        self.mainloop()
    def CheckValid(self):
        status=True
        status_address=True
        status_phone=True
        if len(phone_entry_shipping.get()) == 10 and phone_entry_shipping.get()[0] == "0":
            for i in phone_entry_shipping.get():
                if i in ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]:
                    pass
                else:
                    status_phone=False
                    status=False
        else:
            status_phone=False
            status=False
        if len(address_entry.get()) <=5:
            status_address=False
            status=False
        if status == False:
            info = ""
            if status_address == False:
                if info == "":
                    info += "Địa chỉ quá ngắn"
                else:
                    info += "\nĐịa chỉ quá ngắn"
            if status_phone == False:
                if info == "":
                    info += "Số điện thoại không hợp lệ"
                else:
                    info += "\nSố điện thoại không hợp lệ"
            mb.showerror("Thông báo",info,master=self)
        else:
            global address_shipping, phone_shipping,shipping_status
            shipping_status=True
            address_shipping=address_entry.get()
            phone_shipping=phone_entry_shipping.get()
            self.destroy()
class ChangePassword(tk.Tk):
    def __init__(self,username=None):
        global old_password_entry,new_password_entry
        tk.Tk.__init__(self)
        self.geometry("400x200")
        self.resizable(0, 0)
        self.title("Đổi mật khẩu")
        self.account_username=username
        canvas=Canvas(master=self,width=400,height=200)
        bg_img=PhotoImage(master=canvas,file="change_password.png")
        bg=Label(self,image=bg_img)
        bg.place(x=0,y=0)
        button_img=PhotoImage(master=canvas,file="Save_delivery.png")
        save_button=Button(master=self,relief="flat", overrelief="flat", activebackground="#ffffff", cursor="hand2",
        foreground="#ffffff", background="#ffffff", borderwidth="0",command=self.CheckValid)
        save_button.configure(image=button_img)
        save_button.place(x=134,y=150)
        old_password_entry=Entry(self,font="Montserrat",show="-",relief="flat")
        old_password_entry.place(x=97,y=62,width=283,height=25)
        new_password_entry = Entry(self, font="Montserrat", relief="flat",show="-")
        new_password_entry.place(x=97, y=99, width=283, height=25)
        self.mainloop()
    def CheckValid(self):
        status=False
        excelfile=load_workbook("UsersData.xlsx")
        old_password=old_password_entry.get()
        new_password=new_password_entry.get()
        if self.account_username==None:
            for row in excelfile.worksheets[0]:
                if row[0].value==account_username:
                    if row[1].value==hashlib.sha512(old_password.encode()).hexdigest():
                        status=True
                        row[1].value=hashlib.sha512(new_password.encode()).hexdigest()
                        excelfile.save("UsersData.xlsx")
                        excelfile.close()
                        self.destroy()
                        mb.showinfo("Thông báo","Thay đổi mật khẩu thành công")
                        break
                    else:
                        break
            if status==False:
                mb.showerror("Thông báo","Mật khẩu cũ không hợp lệ")
        else:
            for row in excelfile.worksheets[0]:
                if row[0].value == self.account_username:
                    if row[1].value == hashlib.sha512(old_password.encode()).hexdigest():
                        status = True
                        row[1].value = hashlib.sha512(new_password.encode()).hexdigest()
                        excelfile.save("UsersData.xlsx")
                        excelfile.close()
                        self.destroy()
                        mb.showinfo("Thông báo", "Thay đổi mật khẩu thành công")
                        break
                    else:
                        break
            if status == False:
                mb.showerror("Thông báo", "Mật khẩu cũ không hợp lệ")
# MainSession("Trịnh Đức Hiếu","Admin")
